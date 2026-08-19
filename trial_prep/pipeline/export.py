"""Stage: export the approved report to DOCX and PDF, per design-spec
section 4 (Export) and the "no unreviewed draft mistaken for final work
product" rule in the security section -- every export is watermarked with
its review status.
"""
from datetime import datetime, timezone

from docx import Document as DocxDocument
from docx.shared import Pt, RGBColor
from fpdf import FPDF

from .audit import log_event


def _watermark_text(case_file: dict) -> str:
    review = case_file["review"]
    if review["status"] == "approved":
        return f"ATTORNEY-REVIEWED DRAFT -- Approved by {review['reviewer']} on {review['reviewed_at']}"
    return "AI-ASSISTED DRAFT -- ATTORNEY REVIEW REQUIRED -- NOT FOR COURT USE"


def _safe(text: str) -> str:
    return (text or "").encode("latin-1", errors="replace").decode("latin-1")


def export_docx(case_file: dict, path):
    doc = DocxDocument()
    sections = case_file["report_draft"]["sections"]

    title = doc.add_heading("Trial Preparation Report", level=0)
    watermark = doc.add_paragraph(_watermark_text(case_file))
    watermark.runs[0].bold = True
    watermark.runs[0].font.color.rgb = RGBColor(0xA2, 0x3B, 0x3B)
    doc.add_paragraph(f"Case ID: {case_file['case_id']}  |  Generated: {case_file['report_draft']['generated_at']}")
    doc.add_paragraph(case_file["report_draft"]["disclaimer"]).italic = True

    doc.add_heading("Case Overview", level=1)
    doc.add_paragraph(sections["case_overview"])

    doc.add_heading("Chronological Timeline", level=1)
    table = doc.add_table(rows=1, cols=4)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    hdr[0].text, hdr[1].text, hdr[2].text, hdr[3].text = "Date", "Event", "Source", "Status"
    for e in case_file["timeline"]:
        row = table.add_row().cells
        row[0].text = e.get("date", "")
        row[1].text = e.get("event", "")
        row[2].text = ", ".join(e.get("source_doc_ids", []))
        row[3].text = e.get("status", "")

    doc.add_heading("Evidence Summary", level=1)
    table = doc.add_table(rows=1, cols=5)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    for i, h in enumerate(["Item", "Supports", "Reliability", "Weaknesses", "Source"]):
        hdr[i].text = h
    for e in case_file["evidence_matrix"]:
        row = table.add_row().cells
        row[0].text = e.get("item", "")
        row[1].text = e.get("supports", "")
        row[2].text = f"{e.get('reliability', '')} - {e.get('reliability_reason', '')}"
        row[3].text = e.get("weaknesses", "")
        row[4].text = e.get("source_doc_id", "")

    doc.add_heading("Witness Summary", level=1)
    for w in case_file["witness_summary"]:
        doc.add_heading(f"{w.get('name', 'Unknown')} ({w.get('role', '')})", level=2)
        doc.add_paragraph(w.get("statement_summary", ""))
        if w.get("contradictions"):
            doc.add_paragraph("Contradictions:", style="Intense Quote")
            for c in w["contradictions"]:
                doc.add_paragraph(f"vs {c.get('with')}: {c.get('detail')}", style="List Bullet")
        if w.get("cross_examination_questions"):
            doc.add_paragraph("Suggested cross-examination questions:")
            for q in w["cross_examination_questions"]:
                doc.add_paragraph(q, style="List Number")

    doc.add_heading("Case Strengths", level=1)
    for s in sections["case_strengths"]:
        doc.add_paragraph(s, style="List Bullet")

    doc.add_heading("Case Weaknesses", level=1)
    for s in sections["case_weaknesses"]:
        doc.add_paragraph(s, style="List Bullet")

    doc.add_heading("Missing Documents", level=1)
    for s in sections["missing_documents"]:
        doc.add_paragraph(s, style="List Bullet")

    doc.add_heading("Risks", level=1)
    for s in sections["risks"]:
        doc.add_paragraph(s, style="List Bullet")

    doc.add_heading("Applicable Laws", level=1)
    lr = case_file["legal_research"]
    for p in lr["applicable_provisions"]:
        doc.add_paragraph(f"{p.get('citation')}: {p.get('relevance')}", style="List Bullet")
    for p in lr["precedents"]:
        doc.add_paragraph(f"{p.get('citation')}: {p.get('holding')} -- {p.get('relevance_to_case')}", style="List Bullet")
    if lr["gaps"]:
        doc.add_paragraph("Research gaps:")
        for g in lr["gaps"]:
            doc.add_paragraph(g, style="List Bullet")

    doc.add_heading("Final Trial Brief", level=1)
    doc.add_paragraph(sections["final_trial_brief"])

    doc.save(str(path))
    log_event(case_file["case_id"], "export", "docx_ok", path=str(path))


class ReportPDF(FPDF):
    """multi_cell(w=0, ...) in this fpdf2 version leaves self.x at the right
    margin instead of resetting to the left margin, so the next full-width
    multi_cell call has zero width left to render into. Reset x on every
    call rather than patching each call site.
    """

    def multi_cell(self, w=0, h=None, txt="", *args, **kwargs):
        self.set_x(self.l_margin)
        return super().multi_cell(w, h, txt, *args, **kwargs)


def export_pdf(case_file: dict, path):
    sections = case_file["report_draft"]["sections"]
    pdf = ReportPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()

    pdf.set_font("Helvetica", "B", 16)
    pdf.multi_cell(0, 10, _safe("Trial Preparation Report"))

    pdf.set_font("Helvetica", "B", 10)
    pdf.set_text_color(162, 59, 59)
    pdf.multi_cell(0, 6, _safe(_watermark_text(case_file)))
    pdf.set_text_color(0, 0, 0)

    pdf.set_font("Helvetica", "", 9)
    pdf.multi_cell(0, 5, _safe(f"Case ID: {case_file['case_id']} | Generated: {case_file['report_draft']['generated_at']}"))
    pdf.set_font("Helvetica", "I", 9)
    pdf.multi_cell(0, 5, _safe(case_file["report_draft"]["disclaimer"]))
    pdf.ln(4)

    def heading(text):
        pdf.set_font("Helvetica", "B", 13)
        pdf.ln(3)
        pdf.multi_cell(0, 8, _safe(text))
        pdf.set_font("Helvetica", "", 10)

    def bullets(items):
        for item in items:
            pdf.multi_cell(0, 6, _safe(f"- {item}"))

    heading("Case Overview")
    pdf.multi_cell(0, 6, _safe(sections["case_overview"]))

    heading("Chronological Timeline")
    for e in case_file["timeline"]:
        pdf.multi_cell(0, 6, _safe(f"{e.get('date')} [{e.get('status')}] - {e.get('event')} "
                                    f"(source: {', '.join(e.get('source_doc_ids', []))})"))

    heading("Evidence Summary")
    for e in case_file["evidence_matrix"]:
        pdf.multi_cell(0, 6, _safe(
            f"{e.get('item')} | supports: {e.get('supports')} | reliability: {e.get('reliability')} "
            f"({e.get('reliability_reason')}) | weaknesses: {e.get('weaknesses')}"))

    heading("Witness Summary")
    for w in case_file["witness_summary"]:
        pdf.set_font("Helvetica", "B", 10)
        pdf.multi_cell(0, 6, _safe(f"{w.get('name', 'Unknown')} ({w.get('role', '')})"))
        pdf.set_font("Helvetica", "", 10)
        pdf.multi_cell(0, 6, _safe(w.get("statement_summary", "")))
        for c in w.get("contradictions", []):
            pdf.multi_cell(0, 6, _safe(f"  contradiction vs {c.get('with')}: {c.get('detail')}"))
        for q in w.get("cross_examination_questions", []):
            pdf.multi_cell(0, 6, _safe(f"  Q: {q}"))

    heading("Case Strengths")
    bullets(sections["case_strengths"])
    heading("Case Weaknesses")
    bullets(sections["case_weaknesses"])
    heading("Missing Documents")
    bullets(sections["missing_documents"])
    heading("Risks")
    bullets(sections["risks"])

    heading("Applicable Laws")
    lr = case_file["legal_research"]
    for p in lr["applicable_provisions"]:
        pdf.multi_cell(0, 6, _safe(f"{p.get('citation')}: {p.get('relevance')}"))
    for p in lr["precedents"]:
        pdf.multi_cell(0, 6, _safe(f"{p.get('citation')}: {p.get('holding')} -- {p.get('relevance_to_case')}"))
    if lr["gaps"]:
        pdf.set_font("Helvetica", "I", 9)
        for g in lr["gaps"]:
            pdf.multi_cell(0, 6, _safe(f"Gap: {g}"))
        pdf.set_font("Helvetica", "", 10)

    heading("Final Trial Brief")
    pdf.multi_cell(0, 6, _safe(sections["final_trial_brief"]))

    pdf.output(str(path))
    log_event(case_file["case_id"], "export", "pdf_ok", path=str(path))
